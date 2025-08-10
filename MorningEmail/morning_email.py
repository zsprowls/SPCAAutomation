import pandas as pd
import re
from datetime import datetime, timedelta
import os
from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

def get_user_dates():
    while True:
        try:
            date_input = input("Enter the dates to include (mm/dd/yyyy, separated by commas): ")
            dates = [date.strip() for date in date_input.split(',') if date.strip()]
            # Validate each date
            validated_dates = []
            for date in dates:
                datetime.strptime(date, '%m/%d/%Y')
                validated_dates.append(date)
            return validated_dates
        except ValueError:
            print("Invalid date format. Please use mm/dd/yyyy format and separate dates with commas.")

# Use relative path to the __Load Files Go Here__ directory
LOAD_FILES_DIR = '__Load Files Go Here__'

# Define the stage mappings
STAGE_MAPPINGS = {
    'In Foster': 'In Foster',
    'Hold – SAFE Foster': 'Hold SAFE Foster',
    'Hold - Foster': 'Hold Foster',
    'Hold - Cruelty Foster': 'Hold Cruelty Foster',
    'Hold - Behavior Foster': 'Hold Behavior Foster',
    'Hold - Surgery': 'Hold Surgery',
    'Hold - Doc': 'Hold Doc',
    'Hold - Behavior': 'Hold Behavior',
    'Hold - Dental': 'Hold Dental',
    'Hold - Behavior Mod.': 'Hold Behavior Mod.',
    'Hold - Complaint': 'Hold Complaint',
    'Hold - Stray': 'Hold Stray/Legal',
    'Hold - Legal Notice': 'Hold Stray/Legal'
}

# Define the order for the report
STAGE_ORDER = [
    'In Foster',
    'Hold SAFE Foster',
    'Hold Foster',
    'Hold Cruelty Foster',
    'Hold Behavior Foster',
    'Hold Surgery',
    'Hold Dental',
    'Hold Doc',
    'Hold Behavior',
    'Hold Behavior Mod.',
    'Hold Complaint',
    'Hold Stray/Legal'
]

# Add this list in the order of your mapping
INTAKE_GROUP_ORDER = [
    'Transfer In', 'DOA', 'Euthanasia Request', 'Euthanasia Req – Field', 'Field – Stray', 'Field – OS',
    'Seized – Abandoned', 'Seized – Cruelty', 'Seized – General', 'Seized – Hospital', 'Seized – Signed over',
    'Seized – Eviction', 'Seized – Police', 'Seized – Owner Died', 'Seized – Order Violation', 'Seized - Hoarding',
    'Return', 'Stray', 'OTC – OS', 'OTC - OS - SAFE', 'Clinic - Medical Treatment', 'Clinic - Stray',
    'Clinic - Retention', 'Clinic - Case Assistance', 'Clinic - Case Assistance - Outreach', 'Clinic - Outreach', 'Boarder'
]

def get_fur_fits_count(check_dates):
    # Read FosterCurrent.csv, skipping first 6 rows
    df_foster = pd.read_csv(os.path.join(LOAD_FILES_DIR, 'FosterCurrent.csv'), skiprows=6)
    
    # Convert StartStatusDate to datetime and normalize to remove time component
    df_foster['StartStatusDate'] = pd.to_datetime(df_foster['StartStatusDate'], format='mixed').dt.normalize()
    check_datetimes = [pd.to_datetime(date, format='%m/%d/%Y').normalize() for date in check_dates]
    
    # Count entries where Location is "If The Fur Fits" and StartStatusDate matches any check date
    fur_fits_count = len(df_foster[
        df_foster['FosterReason'].fillna('').str.startswith('Possible Adoption') & 
        (df_foster['StartStatusDate'].isin(check_datetimes))
    ])
    
    return fur_fits_count

def get_foster_count():
    # Read FosterCurrent.csv, skipping first 6 rows
    foster_path = os.path.join(LOAD_FILES_DIR, 'FosterCurrent.csv')
    df_foster = pd.read_csv(foster_path, skiprows=6)
    
    # Filter for rows where Location is either 'Foster Home' or 'If The Fur Fits'
    foster_locations = ['Foster Home', 'If The Fur Fits']
    foster_df = df_foster[df_foster['Location'].isin(foster_locations)]
    
    # Count unique Animal IDs from textbox9
    unique_foster_count = foster_df['textbox9'].nunique()
    
    return unique_foster_count

def get_stage_counts():
    # Read the CSV file, skipping the first 3 rows
    inventory_path = os.path.join(LOAD_FILES_DIR, 'AnimalInventory.csv')
    df = pd.read_csv(inventory_path, skiprows=3)
    
    # Map the stages using our defined mappings
    df['Mapped_Stage'] = df['Stage'].map(STAGE_MAPPINGS)
    
    # Count the animals in each mapped stage
    stage_counts = df['Mapped_Stage'].value_counts()
    
    # Create a dictionary with all stages initialized to 0
    final_counts = {stage: 0 for stage in STAGE_ORDER}
    
    # Update counts for stages that have animals
    for stage, count in stage_counts.items():
        if stage in final_counts:
            final_counts[stage] = count
    
    # Update In Foster count from FosterCurrent.csv
    final_counts['In Foster'] = get_foster_count()
            
    # Return counts in the specified order
    return {stage: final_counts[stage] for stage in STAGE_ORDER}

def get_occupancy_counts():
    # Read the CSV file, skipping first 3 rows
    df = pd.read_csv(os.path.join(LOAD_FILES_DIR, 'AnimalInventory.csv'), skiprows=3)
    
    # Convert DateOfBirth to datetime
    df['DateOfBirth'] = pd.to_datetime(df['DateOfBirth'], format='mixed')
    
    # Get the most recent date from the data to use as "today"
    today = pd.Timestamp.now().normalize()  # Use today's date without the time component
    
    # Calculate age in months - Fix calculation for very young animals
    df['AgeMonths'] = ((today - df['DateOfBirth']).dt.total_seconds() / (60 * 60 * 24 * 30.44))
    
    # Define offsite locations
    offsite_locations = ['Foster Home', 'If The Fur Fits', 'Offsite Adoptions']
    
    # Create location category
    df['LocationCategory'] = df['Location'].apply(
        lambda x: 'Foster/Offsite' if x in offsite_locations else 'Shelter'
    )
    
    # Categorize animals
    def get_category(row):
        animal_type = row['AnimalType'].strip()  # Remove any whitespace
        if animal_type not in ['Cat', 'Dog']:
            return 'Other'  # Bird, Livestock, Equine, and any other types
        elif animal_type == 'Cat':
            return 'Kitten' if row['AgeMonths'] < 6 else 'Cat'
        else:  # Must be Dog
            return 'Puppy' if row['AgeMonths'] < 6 else 'Dog'
    
    df['Category'] = df.apply(get_category, axis=1)
    
    # Create the counts
    categories = ['Cat', 'Dog', 'Kitten', 'Other', 'Puppy']
    counts = {
        'Species/Age': categories + ['TOTAL'],
        'Animals in Shelter': [],
        'Animals in Foster/Off-Site': []
    }
    
    # Get counts for each category
    for category in categories:
        shelter_count = len(df[(df['Category'] == category) & 
                             (df['LocationCategory'] == 'Shelter')])
        offsite_count = len(df[(df['Category'] == category) & 
                              (df['LocationCategory'] == 'Foster/Offsite')])
        
        counts['Animals in Shelter'].append(shelter_count)
        counts['Animals in Foster/Off-Site'].append(offsite_count)
    
    # Add totals
    counts['Animals in Shelter'].append(sum(counts['Animals in Shelter']))
    counts['Animals in Foster/Off-Site'].append(sum(counts['Animals in Foster/Off-Site']))
    
    return pd.DataFrame(counts)

def get_adoptions_count(check_dates):
    # Read AnimalOutcome.csv, skipping the first 3 rows to get to the header
    try:
        df_outcomes = pd.read_csv(os.path.join(LOAD_FILES_DIR, 'AnimalOutcome.csv'), skiprows=3)
    except FileNotFoundError:
        print("AnimalOutcome.csv not found. Returning 0 adoptions.")
        return 0

    # Convert the date column to datetime and extract only the date component
    df_outcomes['Textbox50'] = pd.to_datetime(df_outcomes['Textbox50'], format='mixed').dt.date
    check_dates_dt = [pd.to_datetime(date, format='%m/%d/%Y').date() for date in check_dates]
    
    # Count adoptions for the specified dates
    adoptions_count = len(df_outcomes[
        (df_outcomes['textbox16'] == 'Adoption') & 
        (df_outcomes['Textbox50'].isin(check_dates_dt))
    ])
    
    return adoptions_count

def map_intake_group(row):
    op_type = str(row['OperationType']).strip().upper()
    op_subtype = str(row['OperationSubType']).strip().upper()
    # Logic from mapping
    if op_type == 'TRANSFER IN':
        return 'Transfer In'
    if op_type == 'OWNER/GUARDIAN SURRENDER' and op_subtype == 'DOA':
        return 'DOA'
    if op_type == 'OWNER/GUARDIAN SURRENDER' and op_subtype in ['EUTHANASIA REQUEST', 'EUTHANASIA REQUEST - OTC!']:
        return 'Euthanasia Request'
    if (op_type == 'SEIZED / CUSTODY' and op_subtype == 'SIGNED OVER/EUTHANASIA REQUEST') or \
       (op_type == 'OWNER/GUARDIAN SURRENDER' and op_subtype == 'EUTHANASIA REQUEST - FIELD!'):
        return 'Euthanasia Req – Field'
    if op_type == 'STRAY' and ('FIELD' in op_subtype):
        return 'Field – Stray'
    if op_type == 'OWNER/GUARDIAN SURRENDER' and ('FIELD' in op_subtype) and op_subtype != 'EUTHANASIA REQUEST - FIELD!':
        return 'Field – OS'
    if op_type == 'SEIZED / CUSTODY' and op_subtype == 'ABANDONED':
        return 'Seized – Abandoned'
    if op_type == 'SEIZED / CUSTODY' and op_subtype == 'CRUELTY':
        return 'Seized – Cruelty'
    if op_type == 'SEIZED / CUSTODY' and op_subtype == 'GENERAL':
        return 'Seized – General'
    if op_type == 'SEIZED / CUSTODY' and op_subtype == 'HOSPITAL':
        return 'Seized – Hospital'
    if op_type == 'SEIZED / CUSTODY' and op_subtype == 'SIGNED OVER':
        return 'Seized – Signed over'
    if op_type == 'SEIZED / CUSTODY' and op_subtype == 'EVICTION':
        return 'Seized – Eviction'
    if op_type == 'SEIZED / CUSTODY' and op_subtype == 'POLICE':
        return 'Seized – Police'
    if op_type == 'SEIZED / CUSTODY' and op_subtype == 'OWNER DIED':
        return 'Seized – Owner Died'
    if op_type == 'SEIZED / CUSTODY' and op_subtype == 'COURT ORDER VIOLATION':
        return 'Seized – Order Violation'
    if op_type == 'SEIZED / CUSTODY' and op_subtype == 'HOARDING':
        return 'Seized - Hoarding'
    if op_type == 'RETURN':
        return 'Return'
    if op_type == 'STRAY' and not ('FIELD' in op_subtype):
        return 'Stray'
    if op_type == 'OWNER/GUARDIAN SURRENDER' and ('OTC' in op_subtype):
        return 'OTC – OS'
    if op_type == 'OWNER/GUARDIAN SURRENDER' and op_subtype == 'SAFE FOSTER':
        return 'OTC - OS - SAFE'
    if op_type == 'CLINIC' and op_subtype == 'MEDICAL TREATMENT':
        return 'Clinic - Medical Treatment'
    if op_type == 'CLINIC' and op_subtype == 'STRAY':
        return 'Clinic - Stray'
    if op_type == 'CLINIC' and op_subtype == 'RETENTION':
        return 'Clinic - Retention'
    if op_type == 'CLINIC' and op_subtype == 'CASE ASSISTANCE':
        return 'Clinic - Case Assistance'
    if op_type == 'CLINIC' and op_subtype == 'CASE - OUTREACH':
        return 'Clinic - Case Assistance - Outreach'
    if op_type == 'CLINIC' and op_subtype == 'OUTREACH':
        return 'Clinic - Outreach'
    if op_type == 'BOARDER':
        return 'Boarder'
    return 'not counted'

def get_intake_count_detail(check_dates):
    intake_path = os.path.join(LOAD_FILES_DIR, 'AnimalIntake.csv')
    df_intake = pd.read_csv(intake_path, skiprows=3)
    # Filter by intake date (textbox44)
    df_intake['textbox44'] = pd.to_datetime(df_intake['textbox44'], format='%m/%d/%Y %I:%M %p').dt.date
    check_dates_dt = [pd.to_datetime(date).date() for date in check_dates]
    filtered = df_intake[df_intake['textbox44'].isin(check_dates_dt)].copy()
    # Assign group
    filtered['IntakeGroup'] = filtered.apply(map_intake_group, axis=1)
    # For detail, include all rows (even not counted)
    all_rows = df_intake.copy()
    all_rows['IntakeGroup'] = all_rows.apply(map_intake_group, axis=1)
    # Only keep relevant columns
    detail = all_rows[['AnimalNumber', 'Species', 'textbox44', 'OperationType', 'OperationSubType', 'IntakeGroup']]
    return detail

def get_intake_summary(detail_df):
    # Only use rows that are in the mapping order (including not counted)
    summary = []
    for group in INTAKE_GROUP_ORDER:
        group_df = detail_df[detail_df['IntakeGroup'] == group]
        cat_count = (group_df['Species'].str.lower() == 'cat').sum() if not group_df.empty else 0
        dog_count = (group_df['Species'].str.lower() == 'dog').sum() if not group_df.empty else 0
        other_count = (~group_df['Species'].str.lower().isin(['cat', 'dog'])).sum() if not group_df.empty else 0
        total_count = cat_count + dog_count + other_count
        summary.append({'Group': group, 'Cat': cat_count, 'Dog': dog_count, 'Other': other_count, 'Total': total_count})
    return pd.DataFrame(summary)

def get_hold_stray_data():
    # Add Hold - Stray cases from StageReview.csv
    stageReview_df = pd.read_csv(os.path.join(LOAD_FILES_DIR, 'StageReview.csv'), skiprows=3)
    hold_stray_df = stageReview_df[stageReview_df['Stage'] == 'Hold - Stray']
    
    if not hold_stray_df.empty:
        hold_stray_data = hold_stray_df.apply(
            lambda row: [
                row['textbox89'],  # Animal ID
                f"{row['Location']}, {row['SubLocation']}", # Location info
                pd.to_datetime(row['ReviewDate']).strftime('%Y-%m-%d') if pd.notna(row['ReviewDate']) else 'No Review Date'  # Date in YYYY-MM-DD format
            ], 
            axis=1
        ).tolist()
        return hold_stray_data
    return []

def create_todo_table(doc, category):
    """Create a 2x4 table for Things to Do section"""
    table = doc.add_table(rows=2, cols=4)
    table.style = 'Table Grid'
    
    # Set header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Species'
    header_cells[1].text = 'Name'
    header_cells[2].text = 'Location'
    header_cells[3].text = 'Comments/Concerns'
    
    # Make header bold
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Leave second row blank for manual entry
    for cell in table.rows[1].cells:
        cell.text = ''
    
    return table

def map_outcome_group(row):
    op_type = str(row['OperationType']).strip().upper()
    op_subtype = str(row['OperationSubType']).strip().upper()
    
    # Common outcome groupings - you can adjust these based on your OutcomeGrouping.xlsx
    if op_type == 'ADOPTION':
        if 'OFFSITE' in op_subtype:
            return 'Adoption - Offsite'
        elif 'NEW ADOPTER' in op_subtype:
            return 'Adoption - New Adopter'
        else:
            return 'Adoption - Other'
    elif op_type == 'RETURN TO OWNER/GUARDIAN':
        return 'Return to Owner'
    elif op_type == 'TRANSFER OUT':
        return 'Transfer Out'
    elif op_type == 'CLINIC OUT':
        return 'Clinic Out'
    elif op_type == 'MISSING':
        return 'Missing'
    elif op_type == 'WILDLIFE RELEASE':
        return 'Wildlife Release'
    elif op_type == 'DIED':
        return 'Died'
    elif op_type == 'EUTHANASIA':
        if 'REQUESTED SLEEP' in op_subtype:
            return 'Euthanasia - Requested Sleep'
        elif 'HUMANE REASONS' in op_subtype:
            return 'Euthanasia - Humane Reasons'
        else:
            return 'Euthanasia - Other'
    elif op_type == 'DOA':
        return 'DOA'
    else:
        return 'Other'

def get_outcome_count_detail(check_dates):
    outcome_path = os.path.join(LOAD_FILES_DIR, 'AnimalOutcome.csv')
    df_outcome = pd.read_csv(outcome_path, skiprows=3)
    
    # Filter by outcome date (Textbox50)
    df_outcome['Textbox50'] = pd.to_datetime(df_outcome['Textbox50'], format='mixed').dt.date
    check_dates_dt = [pd.to_datetime(date).date() for date in check_dates]
    filtered = df_outcome[df_outcome['Textbox50'].isin(check_dates_dt)].copy()
    
    # Assign group
    filtered['OutcomeGroup'] = filtered.apply(map_outcome_group, axis=1)
    
    # For detail, include all rows (even not counted)
    all_rows = df_outcome.copy()
    all_rows['OutcomeGroup'] = all_rows.apply(map_outcome_group, axis=1)
    
    # Only keep relevant columns
    detail = all_rows[['AnimalNumber', 'Species', 'Textbox50', 'OperationType', 'OperationSubType', 'OutcomeGroup']]
    return detail

def get_outcome_summary(detail_df):
    # Only show RTOs and Transfer Outs for this specific chart
    # This helps explain inflated intake numbers when animals come in DOA and need to be RTO'd/transferred
    OUTCOME_GROUP_ORDER = [
        'Return to Owner', 'Transfer Out'
    ]
    
    summary = []
    for group in OUTCOME_GROUP_ORDER:
        group_df = detail_df[detail_df['OutcomeGroup'] == group]
        cat_count = (group_df['Species'].str.lower() == 'cat').sum() if not group_df.empty else 0
        dog_count = (group_df['Species'].str.lower() == 'dog').sum() if not group_df.empty else 0
        other_count = (~group_df['Species'].str.lower().isin(['cat', 'dog'])).sum() if not group_df.empty else 0
        total_count = cat_count + dog_count + other_count
        summary.append({'Group': group, 'Cat': cat_count, 'Dog': dog_count, 'Other': other_count, 'Total': total_count})
    return pd.DataFrame(summary)

def export_to_word(check_dates):
    # Get all our data
    adoptions = get_adoptions_count(check_dates)
    fur_fits = get_fur_fits_count(check_dates)
    foster_holds = get_stage_counts()
    occupancy = get_occupancy_counts()
    hold_stray_data = get_hold_stray_data()
    
    # Create Word document
    doc = Document()
    
    # Set all headings to black color
    for style in doc.styles:
        if style.name.startswith('Heading'):
            style.font.color.rgb = None  # This removes any color formatting, making it black
    
    # Get current date for stage and occupancy sections
    current_date = datetime.now().strftime('%m/%d/%y')
    check_dates_str = ', '.join(check_dates)
    
    # Adoptions section - format as date with bold labels and regular numbers
    p = doc.add_paragraph()
    p.add_run(f'{check_dates_str} ').bold = True
    p.add_run('Adoptions: ').bold = True
    p.add_run(f'{adoptions}')
    p.add_run(' & ')
    p.add_run('ITFF: ').bold = True
    p.add_run(f'{fur_fits}')
    
    # Stage Count section
    doc.add_heading(f'Stage Count: {current_date}', level=1)
    
    # Create stage table
    stage_table = doc.add_table(rows=len(foster_holds) + 1, cols=2)
    stage_table.style = 'Table Grid'
    
    # Add header
    stage_table.rows[0].cells[0].text = 'Stage'
    stage_table.rows[0].cells[1].text = 'Count'
    
    # Make header bold
    for cell in stage_table.rows[0].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Add data
    for i, (stage, count) in enumerate(foster_holds.items(), 1):
        stage_table.rows[i].cells[0].text = stage
        stage_table.rows[i].cells[1].text = str(count)
    
    # Add blank line
    doc.add_paragraph()
    
    # Occupancy section
    doc.add_heading(f'Occupancy: {current_date}', level=1)
    
    # Add bullet points under occupancy header
    p1 = doc.add_paragraph()
    p1.add_run('• ').bold = False
    p1.add_run('Animals in the shelter include farm and permanent residents.').bold = False
    
    p2 = doc.add_paragraph()
    p2.add_run('• ').bold = False
    p2.add_run('Animals in Foster/Off-site include in stores and foster care').bold = False
    
    p3 = doc.add_paragraph()
    p3.add_run('• ').bold = False
    p3.add_run('Puppies and Kittens are 5 months and under').bold = False
    
    # Create occupancy table
    occupancy_table = doc.add_table(rows=len(occupancy) + 1, cols=3)
    occupancy_table.style = 'Table Grid'
    
    # Add header
    occupancy_table.rows[0].cells[0].text = 'Species/Age'
    occupancy_table.rows[0].cells[1].text = 'Animals in Shelter'
    occupancy_table.rows[0].cells[2].text = 'Animals in Foster/Off-Site'
    
    # Make header bold
    for cell in occupancy_table.rows[0].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Add data
    for i, row in occupancy.iterrows():
        occupancy_table.rows[i + 1].cells[0].text = str(row['Species/Age'])
        occupancy_table.rows[i + 1].cells[1].text = str(row['Animals in Shelter'])
        occupancy_table.rows[i + 1].cells[2].text = str(row['Animals in Foster/Off-Site'])
    
    # Add blank line
    doc.add_paragraph()
    
    # Things to Do section
    doc.add_heading('Things to Do:', level=1)
    
    # Medical
    doc.add_heading('Medical', level=2)
    create_todo_table(doc, 'Medical')
    doc.add_paragraph()
    
    # Behavior
    doc.add_heading('Behavior', level=2)
    create_todo_table(doc, 'Behavior')
    doc.add_paragraph()
    
    # Legal
    doc.add_heading('Legal', level=2)
    create_todo_table(doc, 'Legal')
    doc.add_paragraph()
    
    # Misc
    doc.add_heading('Misc. - Move me!', level=2)
    create_todo_table(doc, 'Misc')
    doc.add_paragraph()
    
    # Stray section
    doc.add_heading('Stray', level=1)
    
    if hold_stray_data:
        # Create stray table
        stray_table = doc.add_table(rows=len(hold_stray_data) + 1, cols=3)
        stray_table.style = 'Table Grid'
        
        # Add header
        stray_table.rows[0].cells[0].text = 'Animal ID'
        stray_table.rows[0].cells[1].text = 'Location'
        stray_table.rows[0].cells[2].text = 'Review Date'
        
        # Make header bold
        for cell in stray_table.rows[0].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        # Add data
        for i, row_data in enumerate(hold_stray_data, 1):
            stray_table.rows[i].cells[0].text = str(row_data[0])
            stray_table.rows[i].cells[1].text = str(row_data[1])
            stray_table.rows[i].cells[2].text = str(row_data[2])
    else:
        doc.add_paragraph('No stray animals found.')
    
    # Add blank line
    doc.add_paragraph()
    
    # Intake section
    doc.add_heading(f'Intake: {check_dates_str}', level=1)
    doc.add_paragraph('*Please note that most of the animals that come in under subtypes of clinic and stray get worked up and leave the building the same day unless deemed medically necessary.')
    
    # Get intake summary with totals
    intake_detail = get_intake_count_detail(check_dates)
    intake_summary = get_intake_summary(intake_detail)
    
    # Add total row to intake summary
    total_row = {
        'Group': 'TOTAL',
        'Cat': intake_summary['Cat'].sum(),
        'Dog': intake_summary['Dog'].sum(),
        'Other': intake_summary['Other'].sum(),
        'Total': intake_summary['Total'].sum()
    }
    intake_summary = pd.concat([intake_summary, pd.DataFrame([total_row])], ignore_index=True)
    
    # Create intake table
    intake_table = doc.add_table(rows=len(intake_summary) + 1, cols=5)
    intake_table.style = 'Table Grid'
    
    # Add header
    intake_table.rows[0].cells[0].text = 'Group'
    intake_table.rows[0].cells[1].text = 'Cat'
    intake_table.rows[0].cells[2].text = 'Dog'
    intake_table.rows[0].cells[3].text = 'Other'
    intake_table.rows[0].cells[4].text = 'Total'
    
    # Make header bold
    for cell in intake_table.rows[0].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Add data
    for i, row in intake_summary.iterrows():
        intake_table.rows[i + 1].cells[0].text = str(row['Group'])
        intake_table.rows[i + 1].cells[1].text = str(row['Cat']) if row['Cat'] != 0 else ''
        intake_table.rows[i + 1].cells[2].text = str(row['Dog']) if row['Dog'] != 0 else ''
        intake_table.rows[i + 1].cells[3].text = str(row['Other']) if row['Other'] != 0 else ''
        intake_table.rows[i + 1].cells[4].text = str(row['Total']) if row['Total'] != 0 else ''
    
    # Add blank line
    doc.add_paragraph()
    
    # RTOs & Transfers section - shows animals that left through these channels
    # This helps explain inflated intake numbers when animals come in DOA and need to be RTO'd/transferred
    doc.add_heading(f'RTOs & Transfers: {check_dates_str}', level=1)
    
    # Get outcome summary with totals
    outcome_detail = get_outcome_count_detail(check_dates)
    outcome_summary = get_outcome_summary(outcome_detail)
    
    # Add total row to outcome summary
    outcome_total_row = {
        'Group': 'TOTAL',
        'Cat': outcome_summary['Cat'].sum(),
        'Dog': outcome_summary['Dog'].sum(),
        'Other': outcome_summary['Other'].sum(),
        'Total': outcome_summary['Total'].sum()
    }
    outcome_summary = pd.concat([outcome_summary, pd.DataFrame([outcome_total_row])], ignore_index=True)
    
    # Create outcome table
    outcome_table = doc.add_table(rows=len(outcome_summary) + 1, cols=5)
    outcome_table.style = 'Table Grid'
    
    # Add header
    outcome_table.rows[0].cells[0].text = 'Group'
    outcome_table.rows[0].cells[1].text = 'Cat'
    outcome_table.rows[0].cells[2].text = 'Dog'
    outcome_table.rows[0].cells[3].text = 'Other'
    outcome_table.rows[0].cells[4].text = 'Total'
    
    # Make header bold
    for cell in outcome_table.rows[0].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Add data
    for i, row in outcome_summary.iterrows():
        outcome_table.rows[i + 1].cells[0].text = str(row['Group'])
        outcome_table.rows[i + 1].cells[1].text = str(row['Cat']) if row['Cat'] != 0 else ''
        outcome_table.rows[i + 1].cells[2].text = str(row['Dog']) if row['Dog'] != 0 else ''
        outcome_table.rows[i + 1].cells[3].text = str(row['Other']) if row['Other'] != 0 else ''
        outcome_table.rows[i + 1].cells[4].text = str(row['Total']) if row['Total'] != 0 else ''
    
    # Save the document
    output_path = os.path.join('MorningEmail', 'morning_email.docx')
    doc.save(output_path)
    print(f"Word document saved as: {output_path}")

if __name__ == "__main__":
    check_dates = get_user_dates()
    export_to_word(check_dates) 